"""ObservePoint-themed account research dossier (.docx) for research-account.

Input: a scored.json object (the classification object + a `score` block from score_account.py).
This is an INTERNAL AE artifact. The "best opening angle" is labeled internal strategy (the sharp
legal framing belongs to the AE's strategy, never to prospect-facing copy — that is the future
sequence-contacts skill's job, under the tone governor).

Theming mirrors ObservePoint's NERD-app account-detail screen: score badge + status chip header,
color-coded why-now category chips with clickable source hyperlinks, dark section header bars
(card feel), a left-border callout for the opening angle, and green/red verification chips on
contacts. Page stays light/printable.

CLI:  build_dossier.py <scored.json> <out.docx>   (prints the output path)
"""
import json
import pathlib
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Montserrat"
DARK = RGBColor(0x1E, 0x1E, 0x1E)
GRAY = RGBColor(0x5C, 0x5C, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xF3, 0x41, 0x46)
GREEN = RGBColor(0x1F, 0x9D, 0x55)
DARK_HEX, YELLOW_HEX, LIGHT_HEX = "1E1E1E", "F2CD14", "F2F2F2"
RED_HEX, GREEN_HEX, MIDGRAY_HEX, LINK_HEX = "F34146", "1F9D55", "E2E2E2", "0563C1"
LOGO = pathlib.Path(__file__).resolve().parent.parent / "assets" / "op-logo.png"

# why-now category -> (chip fill hex, chip text color). High-severity red, medium yellow, else gray.
_CAT_CHIP = {
    "litigation": (RED_HEX, WHITE), "enforcement": (RED_HEX, WHITE), "incident": (RED_HEX, WHITE),
    "leadership": (YELLOW_HEX, DARK), "hiring": (YELLOW_HEX, DARK), "earnings": (YELLOW_HEX, DARK),
}


def _run(p, text, *, bold=False, size=10.5, color=DARK):
    r = p.add_run(text)
    r.font.name, r.font.bold, r.font.size, r.font.color.rgb = FONT, bold, Pt(size), color
    return r


def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _no_borders(t):
    """Strip default table borders for a clean card/badge look."""
    tblPr = t._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        b.append(e)
    tblPr.append(b)


def _left_accent(cell, hex_color=YELLOW_HEX, sz="24"):
    """A thick colored left border — the NERD 'callout' accent."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = OxmlElement("w:left")
    for k, v in (("w:val", "single"), ("w:sz", sz), ("w:space", "0"), ("w:color", hex_color)):
        el.set(qn(k), v)
    borders.append(el)


def _hyperlink(paragraph, url, text, *, size=9):
    """A real clickable Word hyperlink (blue, underlined)."""
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rPr.append(rf)
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz_el)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), LINK_HEX)
    rPr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    link.append(r)
    paragraph._p.append(link)
    return link


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name, st.font.size, st.font.color.rgb = FONT, Pt(10.5), DARK


def _yellow_bar(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "18"), ("w:space", "1"), ("w:color", YELLOW_HEX)):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def _section(doc, title):
    """Dark 'card header' bar with white title — the NERD panel-header feel. Replaces _heading."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(t)
    c = t.rows[0].cells[0]
    _shade(c, DARK_HEX)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run(p, title.upper(), bold=True, size=11, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _para(doc, text, *, size=10.5, color=DARK, bold=False):
    p = doc.add_paragraph()
    _run(p, text, size=size, color=color, bold=bold)
    return p


def _bullets(doc, items):
    for it in items or []:
        p = doc.add_paragraph(style="List Bullet")
        _run(p, str(it), size=10)


def _chip(cell, text, fill_hex, text_color=WHITE, size=8):
    """Style a table cell as a colored chip."""
    _shade(cell, fill_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size=size, color=text_color)


# ---------- builder ----------
def build_dossier(data):
    score = data.get("score", {})
    research = data.get("research", {})
    scan = data.get("scan", {})

    doc = Document()
    _set_base_style(doc)

    # 1. Header
    if LOGO.exists():
        try:
            doc.add_picture(str(LOGO), width=Inches(2.0))
        except Exception:
            pass
    _run(doc.add_paragraph(), f"Account Research Dossier — {data.get('account', '')}", bold=True, size=20)
    _yellow_bar(doc)
    sub = [b for b in (data.get("date"),
                       ("Prepared by " + data["prepared_by"]) if data.get("prepared_by") else None,
                       data.get("domain")) if b]
    if sub:
        _run(doc.add_paragraph(), "  ·  ".join(sub), size=9, color=GRAY)

    # 2. Verdict block (score badge + status chip + scores breakdown)
    vt = doc.add_table(rows=1, cols=3)
    _no_borders(vt)
    cells = vt.rows[0].cells
    # col0: score badge
    _chip(cells[0], str(score.get("finalScore", 0)), YELLOW_HEX, DARK, size=26)
    cells[0].width = Inches(0.9)
    # col1: qualified chip
    qualified = score.get("qualified")
    _chip(cells[1],
          "QUALIFIED" if qualified else "NOT QUALIFIED",
          GREEN_HEX if qualified else MIDGRAY_HEX,
          WHITE if qualified else DARK,
          size=11)
    cells[1].width = Inches(1.7)
    # col2: scores breakdown (plain, no chip)
    cells[2].width = Inches(4.0)
    p2 = cells[2].paragraphs[0]
    _run(p2, f"fit {score.get('fitScore', 0)} · why-now {score.get('whyNowScore', 0)}", size=10, color=GRAY)
    if score.get("lowFitHighTrigger"):
        p3 = cells[2].add_paragraph()
        _run(p3, "Qualified on the trigger override despite sub-gate fit — a timing play.", size=9, color=RED)
    if data.get("rationale"):
        _para(doc, data["rationale"], size=10, color=GRAY)

    # 3. Why now
    _section(doc, "Why now")
    # `or 0` guards against a null/absent points value (dict.get's default only covers a MISSING key,
    # not a key mapped to None) so the sort key below never compares None to int.
    pts_by_desc = {b.get("description"): (b.get("points") or 0) for b in score.get("whyNowBreakdown", [])}
    trigs = sorted(data.get("triggers", []) or [],
                   key=lambda t: pts_by_desc.get(t.get("description"), 0), reverse=True)
    if trigs:
        trig_t = doc.add_table(rows=1, cols=5)
        trig_t.style = "Table Grid"
        trig_t.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        headers = ["", "Trigger", "Date", "Points", "Source"]
        for i, h in enumerate(headers):
            hc = trig_t.rows[0].cells[i]
            _shade(hc, DARK_HEX)
            _run(hc.paragraphs[0], h, bold=True, size=9, color=WHITE)
        # Data rows
        for ri, trig in enumerate(trigs):
            row_cells = trig_t.add_row().cells
            # col0: category chip
            category = (trig.get("category") or "").lower()
            (fill, tcolor) = _CAT_CHIP.get(category, (MIDGRAY_HEX, DARK))
            _chip(row_cells[0], (trig.get("category") or "—").upper(), fill, tcolor)
            # cols 1-4: zebra shading on odd rows (not on chip cell col0)
            if ri % 2 == 1:
                for ci in (1, 2, 3, 4):
                    _shade(row_cells[ci], LIGHT_HEX)
            # col1: description
            _run(row_cells[1].paragraphs[0], trig.get("description", ""), size=9)
            # col2: date
            _run(row_cells[2].paragraphs[0], trig.get("date", "—"), size=9)
            # col3: points
            _run(row_cells[3].paragraphs[0], str(pts_by_desc.get(trig.get("description"), 0)), size=9)
            # col4: source hyperlink or dash
            if trig.get("sourceUrl"):
                _hyperlink(row_cells[4].paragraphs[0], trig["sourceUrl"], "source")
            else:
                _run(row_cells[4].paragraphs[0], "—", size=9)
    else:
        _para(doc, "No acute web-tracking trigger event found. A strong fit with no trigger is a "
                   "valid, honest result.", size=10, color=GRAY)

    # 4. ICP fit
    _section(doc, "ICP fit")
    fit_t = doc.add_table(rows=1, cols=4)
    fit_t.style = "Table Grid"
    fit_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    fit_headers = ["Criterion", "Met?", "Points", "Evidence"]
    for i, h in enumerate(fit_headers):
        hc = fit_t.rows[0].cells[i]
        _shade(hc, DARK_HEX)
        _run(hc.paragraphs[0], h, bold=True, size=9, color=WHITE)
    for b in score.get("fitBreakdown", []):
        row_cells = fit_t.add_row().cells
        _run(row_cells[0].paragraphs[0], b["label"], size=9)
        if b["met"]:
            _chip(row_cells[1], "✓ MET", GREEN_HEX, WHITE)
        else:
            _chip(row_cells[1], "—", MIDGRAY_HEX, DARK)
        _run(row_cells[2].paragraphs[0], str(b["points"]), bold=True, size=9)
        _run(row_cells[3].paragraphs[0], b.get("evidence") or "—", size=9)

    # 5. Account overview
    _section(doc, "Account overview")
    if research.get("companyOverview"):
        _para(doc, research["companyOverview"])
    if research.get("painHypotheses"):
        _para(doc, "Why ObservePoint matters here:", bold=True, size=10)
        _bullets(doc, research["painHypotheses"])
    if research.get("competitorIntel"):
        _para(doc, "Competitor intel: " + research["competitorIntel"], size=10)
    # Tech stack + the measured scan inventory
    tech = research.get("techStackNotes", "")
    tags = ", ".join(scan.get("tags") or [])
    cmp_line = scan.get("cmp")
    measured = []
    if cmp_line:
        measured.append(f"CMP: {cmp_line}" + (" (ObservePoint-supported)" if scan.get("cmp_supported") else ""))
    if tags:
        measured.append(f"Tags/pixels: {tags}")
    if scan.get("site_census"):
        measured.append(f"Site Census page count: {scan['site_census']}")
    line = "Tech stack: " + tech
    if measured:
        line += "  |  Measured on-site: " + "; ".join(measured) + "."
    _para(doc, line, size=10)

    # 6. Best opening angle — callout with yellow left accent
    _section(doc, "Best opening angle")
    callout_t = doc.add_table(rows=1, cols=1)
    _no_borders(callout_t)
    callout_cell = callout_t.rows[0].cells[0]
    _shade(callout_cell, LIGHT_HEX)
    _left_accent(callout_cell, YELLOW_HEX)
    p_warn = callout_cell.paragraphs[0]
    _run(p_warn, "Internal strategy — not prospect-facing copy.", bold=True, size=9, color=RED)
    p_angle = callout_cell.add_paragraph()
    _run(p_angle, research.get("bestOpeningAngle", ""), size=10)

    # 7. Contacts
    _section(doc, "Contacts")
    held = 0
    contact_rows_exist = False
    contacts = data.get("contacts", []) or []
    if contacts:
        contact_t = doc.add_table(rows=1, cols=6)
        contact_t.style = "Table Grid"
        contact_t.alignment = WD_TABLE_ALIGNMENT.LEFT
        contact_headers = ["Name", "Title", "LinkedIn", "Verified?", "Hook", "Avoid"]
        for i, h in enumerate(contact_headers):
            hc = contact_t.rows[0].cells[i]
            _shade(hc, DARK_HEX)
            _run(hc.paragraphs[0], h, bold=True, size=9, color=WHITE)
        for c in contacts:
            verified = bool(c.get("sourceVerified")) and bool(c.get("sourceUrl"))
            if not verified:
                held += 1
            row_cells = contact_t.add_row().cells
            _run(row_cells[0].paragraphs[0], c.get("name", ""), size=9)
            _run(row_cells[1].paragraphs[0], c.get("title", ""), size=9)
            _run(row_cells[2].paragraphs[0], c.get("linkedin") or "—", size=9)
            if verified:
                _chip(row_cells[3], "✓ VERIFIED", GREEN_HEX, WHITE)
            else:
                _chip(row_cells[3], "⚠ HELD BACK", RED_HEX, WHITE)
            _run(row_cells[4].paragraphs[0], c.get("personalizationHook", ""), size=9)
            _run(row_cells[5].paragraphs[0], c.get("avoid", ""), size=9)
        contact_rows_exist = True
    if held:
        _para(doc, f"{held} contact(s) held back: missing source verification. Confirm the person and "
                   f"current title before any outreach (no fabricated or unverified contacts ship).",
              size=9, color=RED)

    # 8. Sources & method
    _section(doc, "Sources & method")
    for source_url in research.get("researchSources", []) or []:
        p = doc.add_paragraph()
        _hyperlink(p, source_url, source_url)
    _para(doc, "Method: public web research + an ObservePoint CMP/tag scan of the live site. "
               "The score is computed deterministically from ObservePoint's ICP weights "
               "(reproducible; not a model guess).", size=9, color=GRAY)

    return doc


def main(argv):
    if len(argv) < 3:
        sys.exit("usage: build_dossier.py <scored.json> <out.docx>")
    data = json.loads(pathlib.Path(argv[1]).read_text())
    build_dossier(data).save(argv[2])
    print(argv[2])


if __name__ == "__main__":
    main(sys.argv)
