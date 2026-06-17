"""Customer-facing scope & investment proposal (.docx). Clean by construction — internal
derivation lives in the separate rep-only file (build_internal_evidence.py).

All agent-composed customer-facing strings (monitoring_summary, properties_note, cadence layer
names and "why" lines) are guarded by customer_clean.assert_clean before the document is built.
Identity/factual fields (customer, domains, prepared_by, regulations) are NOT passed to the
guard — see customer_clean caller contract.

Theming matches ObservePoint's brand: Montserrat, near-black #1E1E1E, brand yellow #F2CD14,
light gray #F2F2F2 table rows, and the OP logo in the header.

Input schema (all keys tolerant; orchestrator assembles from derive-page-count + size-and-price):
{
  customer, prepared_by?, date?, use_case, domains[], properties_note?, regulations[],
  monitoring_summary,                       # customer-facing prose (guarded)
  page_count: {low, anchor, high,
               url_total?, defensible?, discounted?, census_id?, crawl_status?, spiral_note?,
               confidence?},               # internal-only keys present but ignored here
  consent_states: {count, names[]},
  multipliers: {geographies, scenarios, environments},  # all three multipliers; geo/env rows shown only when >1
  cadence_layers: [{name, why, pct, runs_per_year}],    # pct required; pages/runs ignored (derived here)
  buffer_pct?,                              # additive buffer fraction (e.g. 0.15) — one pass over combined
  usage: {combined_pages, predicted_scans}, # combined_pages = anchor × multipliers; predicted_scans = Σ(combined×pct×runs)+round(combined×buffer%)
  pricing: {predicted_price, range_low_price, range_high_price,
            price_by_band?: [{band_limit, rate, pages, cost}], pricing_source?,
            modeled_scans?, modeled_price?},
  internal: {…}                            # present but ignored — goes to internal-evidence file
}
"""
import customer_clean
import json
import math
import pathlib
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Montserrat"
DARK = RGBColor(0x1E, 0x1E, 0x1E)
GRAY = RGBColor(0x5C, 0x5C, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_HEX, YELLOW_HEX, LIGHT_HEX = "1E1E1E", "F2CD14", "F2F2F2"
LOGO = pathlib.Path(__file__).resolve().parent.parent / "assets" / "op-logo.png"

_FREQ = {1: "Annually", 4: "Quarterly", 12: "Monthly", 26: "Bi-weekly", 52: "Weekly", 365: "Daily"}


# ---------- formatting helpers ----------
def _int(n):
    return f"{int(round(n)):,}"


def _round_sig(n, sig=2):
    """Round to ~2 significant figures for customer-facing display (matches the site-census
    methodology). Round-to-nearest-1000 erased small footprints (80 -> 0) and collapsed neighbours
    (4,722 & 5,398 -> 5,000); 2-sig-fig keeps 80 -> 80, 4,722 -> 4,700, 5,398 -> 5,400, 95,721 -> 96,000."""
    n = float(n)
    if n == 0:
        return 0
    d = sig - 1 - math.floor(math.log10(abs(n)))
    return int(round(n, d))


def _usd(n):
    return f"${n:,.0f}"


def _run(p, text, *, bold=False, size=10.5, color=DARK):
    r = p.add_run(text)
    r.font.name, r.font.bold, r.font.size, r.font.color.rgb = FONT, bold, Pt(size), color
    return r


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _no_borders(t):
    """Strip default table borders for a clean card/badge look."""
    tblPr = t._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        b.append(e)
    tblPr.append(b)


def _left_accent(cell, hex_color=YELLOW_HEX, sz="24"):
    """A thick colored left border — the NERD 'callout' accent."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = OxmlElement("w:left")
    for k, v in (("w:val", "single"), ("w:sz", sz), ("w:space", "0"), ("w:color", hex_color)):
        el.set(qn(k), v)
    borders.append(el)


def _chip(cell, text, fill_hex, text_color=WHITE, size=8):
    """Style a table cell as a colored chip."""
    _shade(cell, fill_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size=size, color=text_color)


def _yellow_bar(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "18"), ("w:space", "1"), ("w:color", YELLOW_HEX)):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def _heading(doc, text, *, color=DARK, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    _run(p, text, bold=True, size=size, color=color)
    _yellow_bar(doc)


def _section(doc, title, *, fill=DARK_HEX, text_color=WHITE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(t)
    c = t.rows[0].cells[0]
    _shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run(p, title, bold=True, size=12, color=text_color)   # NOTE: preserve case — do NOT .upper()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _para(doc, text, *, size=10.5, color=DARK, bold=False):
    p = doc.add_paragraph()
    _run(p, text, size=size, color=color, bold=bold)
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, DARK_HEX)
        _run(c.paragraphs[0], h, bold=True, size=9, color=WHITE)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[i], LIGHT_HEX)
            bold = str(val).startswith("**")
            _run(cells[i].paragraphs[0], str(val).strip("*"), size=9, bold=bold)
    return t


def _highlight(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    _shade(c, YELLOW_HEX)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size=15, color=DARK)
    return t


def _set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name, st.font.size, st.font.color.rgb = FONT, Pt(10.5), DARK


# ---------- guard ----------
def _assert_clean(data):
    """Guard all agent-composed customer-facing text against internal-only language. Identity/
    factual fields (customer, domains, prepared_by, regulations) are NOT scrubbed — see
    customer_clean caller contract."""
    strings = [data.get("monitoring_summary", ""), data.get("properties_note", "")]
    for L in data.get("cadence_layers") or []:
        strings.append(L.get("name", ""))
        strings.append(L.get("why", ""))
    customer_clean.assert_clean(strings, where="proposal")


# ---------- builder ----------
def build_proposal(data):
    _assert_clean(data)
    pc = data["page_count"]
    pr = data["pricing"]
    us = data["usage"]
    cs = data.get("consent_states", {"count": 1, "names": ["Default"]})
    mx = data.get("multipliers", {"geographies": 1, "scenarios": cs.get("count", 1), "environments": 1})

    err = _sweep_reconcile_error(data)
    if err:
        raise ValueError(err)

    doc = Document()
    _set_base_style(doc)

    # Header: logo + title
    if LOGO.exists():
        try:
            doc.add_picture(str(LOGO), width=Inches(2.0))
        except Exception:
            pass
    _run(doc.add_paragraph(), "Scope & Investment Proposal", bold=True, size=22)
    _yellow_bar(doc)
    _run(doc.add_paragraph(), f"Prepared for {data.get('customer', '')}", size=12, bold=True)
    sub_bits = [b for b in (data.get("date"),
                            ("Prepared by " + data["prepared_by"]) if data.get("prepared_by") else None,
                            (data.get("use_case", "").title() if data.get("use_case") else None)) if b]
    if sub_bits:
        _run(doc.add_paragraph(), "  ·  ".join(sub_bits), size=9, color=GRAY)

    # §1 Footprint
    _section(doc, "1. Your web footprint")
    _para(doc, f"We scanned your web properties to establish how many real pages ObservePoint "
               f"would monitor. Your estimated footprint is approximately "
               f"{_int(_round_sig(pc['anchor']))} pages "
               f"(range {_int(_round_sig(pc['low']))}–{_int(_round_sig(pc['high']))}).")
    if data.get("properties_note"):
        _para(doc, data["properties_note"] + " The full property list is in the attached Scope of "
                   "Work workbook — please review and confirm which properties are in scope.")
    _para(doc, "The attached Scope of Work workbook is a live calculator — adjust the highlighted "
               "inputs and your page-scans and annual investment update automatically. It also lists "
               "every property's page count (Scope Detail) and real example pages (Sample pages).",
          size=9.5, color=GRAY)
    # Footprint badge (page count only — confidence is rep-only, in the internal file)
    badge_t = doc.add_table(rows=1, cols=1)
    _no_borders(badge_t)
    _chip(badge_t.rows[0].cells[0], _int(_round_sig(pc["anchor"])) + " pages", YELLOW_HEX, DARK, size=14)

    # §2 What we monitor
    _section(doc, "2. What ObservePoint will monitor")
    _para(doc, data.get("monitoring_summary", ""))
    _para(doc, "Delivered through ObservePoint Web Audits (automated page scanning), Tag & Variable "
               "Rules (validation of what fires and what must not), and scheduled re-scans — "
               "continuous evidence that your site behaves the way it should.")

    # §3 Derivation
    _section(doc, "3. How your annual usage is calculated")
    _para(doc, "A page scan = one page checked one time. Scanning your pages once is one pass; "
               "checking them on a recurring schedule multiplies that into annual usage.", size=9.5, color=GRAY)
    combined = us.get("combined_pages", us.get("pages_per_sweep"))
    sweep_rows = [
        ["Total Pages Found on your properties", "", _int(pc["anchor"])],
    ]
    if mx.get("geographies", 1) > 1:
        sweep_rows.append([f"× Geographies monitored", "", f"×{mx['geographies']}"])
    sweep_rows.append([f"× Consent states monitored ({', '.join(cs['names'])})", "", f"×{mx['scenarios']}"])
    if mx.get("environments", 1) not in (1, None):
        sweep_rows.append(["× Environments (prod + staging)", "", f"×{mx['environments']}"])
    sweep_rows.append(["**Combined pages monitored**", "", "**" + _int(combined) + "**"])
    _table(doc, ["From pages to combined scope", "", "Pages"], sweep_rows)
    _para(doc, "")
    cadence_rows = []
    for L in data.get("cadence_layers") or []:
        freq = _FREQ.get(L.get("runs_per_year"), f"{L.get('runs_per_year')}×/yr")
        pages_each = round(combined * L["pct"])
        scans = round(combined * L["pct"] * L["runs_per_year"])
        cadence_rows.append([L["name"], L.get("why", ""), freq, _int(pages_each),
                             str(L.get("runs_per_year", "")), _int(scans)])
    # Additive buffer — one pass over combined*buffer%, NOT multiplied by a cadence.
    buffer_pct = data.get("buffer_pct")
    if buffer_pct:
        buffer_scans = round(combined * buffer_pct)
        cadence_rows.append(["Buffer",
                             "Headroom for new pages, campaigns, and ad-hoc re-scans.",
                             "One pass", _int(buffer_scans), "1", _int(buffer_scans)])
    predicted = us.get("predicted_scans", us.get("annual_scans"))
    cadence_rows.append(["**Total annual page scans (predicted)**", "", "", "", "",
                         "**" + _int(predicted) + "**"])
    _table(doc, ["What's monitored", "Why", "How often", "Pages each run", "Runs/yr", "Page scans/yr"],
           cadence_rows)

    # §4 Investment
    _section(doc, "4. Annual investment")
    predicted = us.get("predicted_scans", us.get("annual_scans"))
    predicted_price = pr.get("predicted_price", pr.get("recommended_price"))
    _highlight(doc, f"{_int(predicted)} page scans   ·   {_usd(predicted_price)} / year")
    # NERD callout for the pricing note
    callout_t = doc.add_table(rows=1, cols=1)
    _no_borders(callout_t)
    callout_cell = callout_t.rows[0].cells[0]
    _shade(callout_cell, LIGHT_HEX)
    _left_accent(callout_cell, YELLOW_HEX)
    p_note = callout_cell.paragraphs[0]
    _run(p_note, "Usage-based pricing at ObservePoint's published graduated rates. The annual "
                 "investment is the exact price of your predicted page scans — the cadence above, "
                 "plus buffer headroom.", size=9.5, color=GRAY)
    if pr.get("range_low_price") and pr.get("range_high_price"):
        _para(doc, f"As your property list and monitoring cadence are confirmed, expect a range of "
                   f"{_usd(pr['range_low_price'])}–{_usd(pr['range_high_price'])} per year.")

    # §5 To finalize
    _section(doc, "5. To finalize")
    for item in ("Confirm the in-scope property list (see the attached Scope of Work workbook).",
                 f"Confirm applicable regulations and consent states ({', '.join(data.get('regulations', []) or ['TBD'])}).",
                 "Confirm the monitoring cadence above matches your risk tolerance.",
                 "Finalize the agreement at the confirmed usage."):
        p = doc.add_paragraph(style="List Bullet")
        _run(p, item)

    return doc


_DOC_REF = "see references/deliverables-mapping.md"
# Required top-level key → short expected-shape hint, surfaced in the friendly error.
_REQUIRED = {
    "page_count": "{low, anchor, high}",
    "pricing": "{predicted_price}",
    "usage": "{combined_pages, predicted_scans}",
    "multipliers": "{geographies, scenarios, environments}",
}


def _sweep_reconcile_error(data):
    """Return an actionable message if the §3 multiplier chain does NOT reconcile to
    usage.pages_per_sweep, else None.

    The §3 sweep table walks anchor pages → × geographies → × consent states → × environments
    → 'pages per full sweep'. Those factors MUST multiply the anchor up to the same
    pages_per_sweep the engine (compute_scope) computed. If a factor is dropped from the
    hand-assembled proposal payload (the production bug: geographies forgotten while the engine
    used geos=3), the table silently omits a row and the chain no longer adds up. This guard
    refuses to render a non-reconciling proposal — the §3 factors are not allowed to drift from
    compute_scope's authoritative output. Returns None when other validation owns the failure
    (missing blocks)."""
    pc = data.get("page_count") or {}
    us = data.get("usage") or {}
    cs = data.get("consent_states") or {}
    mx = data.get("multipliers") or {}
    combined = us.get("combined_pages", us.get("pages_per_sweep"))
    anchor = pc.get("anchor")
    if not isinstance(anchor, (int, float)) or not isinstance(combined, (int, float)):
        return None  # _REQUIRED / page_count validation handles missing/garbage blocks
    geos = mx.get("geographies", 1) or 1
    scen = mx.get("scenarios", cs.get("count", 1)) or 1
    env = mx.get("environments", 1) or 1
    product = anchor * geos * scen * env
    tol = max(1, round(0.01 * combined))   # absorb fractional-environment rounding; a dropped factor is ≫1%
    if abs(round(product) - round(combined)) > tol:
        return (f"scope inputs don't reconcile — multipliers "
                f"(geographies×scenarios×environments = {geos}×{scen}×{env}) on {_int(anchor)} "
                f"anchor pages give {_int(product)}, but usage.combined_pages is {_int(combined)}. "
                f"A factor (often geographies) was dropped from the proposal payload — copy "
                f"compute_scope's emitted 'multipliers' into proposal.json verbatim; {_DOC_REF}")
    # Additive-buffer predicted-scans identity: sum(layer combined*pct*runs) + round(combined*buffer%)
    # MUST equal the engine's emitted predicted total. A dropped layer or buffer silently undercounts.
    predicted = us.get("predicted_scans", us.get("annual_scans"))
    if isinstance(predicted, (int, float)):
        layers = data.get("cadence_layers") or []
        layers_sum = sum(round(combined * L["pct"] * L["runs_per_year"]) for L in layers
                         if isinstance(L.get("pct"), (int, float))
                         and isinstance(L.get("runs_per_year"), (int, float)))
        buffer_pct = data.get("buffer_pct") or 0
        recomputed = layers_sum + round(combined * buffer_pct)
        ptol = max(1, round(0.01 * predicted))
        if abs(recomputed - round(predicted)) > ptol:
            return (f"predicted scans don't reconcile — cadence layers "
                    f"(Σ combined×pct×runs = {_int(layers_sum)}) plus buffer "
                    f"(combined×{buffer_pct} = {_int(round(combined * buffer_pct))}) give "
                    f"{_int(recomputed)}, but usage.predicted_scans is {_int(predicted)}. "
                    f"A cadence layer or the buffer was dropped/edited in the proposal payload — "
                    f"copy compute_scope's emitted 'cadence_layers' + 'buffer_pct' verbatim; {_DOC_REF}")
    return None


def _validate(data):
    """Friendly up-front validation so a malformed proposal JSON yields a one-line actionable
    message instead of a raw KeyError traceback a rep can't decode."""
    if not isinstance(data, dict):
        sys.exit(f"scope-calculator: malformed proposal inputs — expected a JSON object; {_DOC_REF}")
    for key, shape in _REQUIRED.items():
        if key not in data or data[key] in (None, {}, []):
            sys.exit(f"scope-calculator: missing/malformed '{key}' — expected {shape}; {_DOC_REF}")
    pc = data["page_count"]
    if not isinstance(pc, dict) or any(k not in pc for k in ("low", "anchor", "high")):
        sys.exit("scope-calculator: missing/malformed 'page_count' — expected "
                 f"{_REQUIRED['page_count']}; {_DOC_REF}")
    err = _sweep_reconcile_error(data)
    if err:
        sys.exit("scope-calculator: " + err)


def main(argv):
    if len(argv) > 1:
        with open(argv[1]) as fh:
            raw = fh.read()
    else:
        raw = sys.stdin.read()
    out = argv[2] if len(argv) > 2 else "proposal.docx"
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        sys.exit(f"scope-calculator: proposal inputs are not valid JSON ({e}); {_DOC_REF}")
    _validate(data)
    build_proposal(data).save(out)  # narrative guard runs inside
    print(out)


if __name__ == "__main__":
    main(sys.argv)
