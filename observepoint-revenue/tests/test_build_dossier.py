import json
import pathlib
import subprocess
import sys

from docx import Document

import build_dossier as bd
import score_account as sa

ROOT = pathlib.Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "skills" / "research-account" / "scripts" / "build_dossier.py"
AS_OF = "2026-06-07"

CLASSIFICATION = {
    "account": "Acme Health", "domain": "acmehealth.com", "prepared_by": "Jarrod Wilbur", "date": AS_OF,
    "scan": {"cmp": "OneTrust", "cmp_supported": True, "cmp_confirmed": True,
             "tags": ["Google Tag Manager", "GA4", "Meta Pixel"],
             "tag_method": "homepage script-signature scan (WebFetch)", "site_census": None,
             "method": "observepoint detect_cmp + static tag scan"},
    "fit": [
        {"key": "privacyConsentSurface", "met": True, "evidence": "OneTrust CMP (confirmed via ObservePoint scan)."},
        {"key": "regulatoryExposure", "met": True, "evidence": "HIPAA + CCPA."},
        {"key": "tagPixelDensity", "met": True, "evidence": "GTM + Meta Pixel observed."},
        {"key": "webScale", "met": True, "evidence": "Thousands of pages."},
        {"key": "targetVertical", "met": True, "evidence": "Healthcare."},
        {"key": "analyticsAccuracy", "met": False},
    ],
    "triggers": [
        {"description": "CIPA class action re: tracking pixels.", "date": "2026-01-16",
         "sourceUrl": "https://example.com/cipa", "category": "litigation", "scoreKey": "pixelWiretapSuit"},
    ],
    "rationale": "Strong privacy fit with active CIPA exposure.",
    "research": {
        "companyOverview": "Acme Health is a large consumer healthcare web property.",
        "keyTriggers": ["Active CIPA class action."],
        "painHypotheses": ["Needs dated per-page evidence of what fires and the consent state."],
        "competitorIntel": "Likely OneTrust CMP; no incumbent pixel-audit vendor.",
        "techStackNotes": "GTM tag orchestration; multiple ad/analytics partners.",
        "bestOpeningAngle": "Offer a free scan tear sheet of what fires vs captured consent.",
        "researchSources": ["https://example.com/leadership", "https://example.com/news"],
    },
    "contacts": [
        {"name": "Dana Rivera", "title": "Senior Director, Privacy Counsel",
         "linkedin": "https://www.linkedin.com/in/example-dana", "sourceVerified": True,
         "sourceUrl": "https://www.linkedin.com/in/example-dana",
         "personalizationHook": "Posted about expanding the privacy team.",
         "toneGuidance": "Privacy-counsel voice.", "avoid": "Do not lead with their lawsuit."},
        {"name": "Sam Okonkwo", "title": "VP, Compliance",
         "linkedin": None, "sourceVerified": False, "sourceUrl": "",
         "personalizationHook": "Built the internal audit program.",
         "toneGuidance": "Operator voice.", "avoid": "Avoid legal-threat framing."},
    ],
}

SCORED = sa.score(CLASSIFICATION, AS_OF)


def _text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
                # Also recurse into nested tables (chips live inside table cells)
                for nested in c.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            parts.append(nc.text)
    return "\n".join(parts)


def _hyperlink_targets(doc):
    return [rel.target_ref for rel in doc.part.rels.values()
            if "hyperlink" in rel.reltype]


def test_header_and_verdict():
    t = _text(bd.build_dossier(SCORED))
    assert "Acme Health" in t
    assert "Account Research Dossier" in t
    assert "QUALIFIED" in t
    assert "120" in t            # final score (fit 90 + why-now 30)
    assert "90" in t and "30" in t  # fit + why-now


def test_why_now_and_fit_sections():
    t = _text(bd.build_dossier(SCORED))
    assert "WHY NOW" in t.upper()
    assert "CIPA class action" in t
    # Source URL is now a real hyperlink, not plain paragraph text — checked separately
    assert "Privacy & consent surface" in t        # a fit-criterion label
    assert "confirmed via ObservePoint scan" in t  # measured evidence folded in
    assert "MET" in t                              # ICP fit chip


def test_source_is_a_clickable_hyperlink():
    doc = bd.build_dossier(SCORED)
    assert "https://example.com/cipa" in _hyperlink_targets(doc)


def test_overview_and_internal_strategy_label():
    t = _text(bd.build_dossier(SCORED))
    assert "consumer healthcare web property" in t
    assert "tear sheet" in t                       # best opening angle text
    assert "Internal strategy" in t                # the strategy-vs-copy label
    assert "Meta Pixel" in t                       # measured tag inventory surfaced


def test_contacts_table_and_held_back_gate():
    t = _text(bd.build_dossier(SCORED))
    assert "Dana Rivera" in t and "Sam Okonkwo" in t
    assert "HELD BACK" in t.upper()               # Sam is unverified -> flagged, not hidden
    assert "VERIFIED" in t.upper()                # Dana is verified -> verified chip


def test_not_qualified_verdict():
    data = {**SCORED, "score": {**SCORED["score"], "qualified": False}}
    t = _text(bd.build_dossier(data))
    assert "NOT QUALIFIED" in t


def test_empty_triggers_renders_graceful_note():
    data = {**SCORED, "triggers": [], "score": {**SCORED["score"], "whyNowBreakdown": []}}
    t = _text(bd.build_dossier(data))
    assert "No acute web-tracking trigger event found" in t


def test_cli_writes_docx(tmp_path):
    f = tmp_path / "scored.json"; f.write_text(json.dumps(SCORED))
    out = tmp_path / "dossier.docx"
    res = subprocess.run([sys.executable, str(SCRIPT), str(f), str(out)],
                         capture_output=True, text=True)
    assert res.returncode == 0, res.stderr
    assert res.stdout.strip() == str(out)
    assert "Acme Health" in _text(Document(out))
