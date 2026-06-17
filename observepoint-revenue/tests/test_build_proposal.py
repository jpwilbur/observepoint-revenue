import copy
import json
import pathlib
import subprocess
import sys

import pytest
from docx import Document

import build_proposal as bp
import compute_scope as cs

ROOT = pathlib.Path(__file__).resolve().parent.parent
SCRIPT = ROOT / "skills" / "scope-calculator" / "scripts" / "build_proposal.py"

# ---------------------------------------------------------------------------
# Step-1 recomputed numbers (engine is authoritative; never hand-arithmetic)
# Regenerate: cs.compute({"page_count":{...anchor...}, "multipliers":{...}, "cadence_layers":_CADENCE_LAYERS,
#   "buffer_pct":0.15, "tiers":cs.BAKED_TIERS})["anchor"] → combined_pages / predicted_scans / price["total"].
# GALLAGHER: anchor=95721, geos=1, scen=3, env=1 → combined=287163
#   Baseline inventory: pages_each=287163, scans=287163
#   High Priority:      pages_each=4307,   scans=223987
#   Moderate Priority:  pages_each=21537,  scans=258447
#   Low Priority:       pages_each=57433,  scans=229730
#   buffer_scans=43074  predicted=1042402  price=97984.12
# GILEAD: anchor=848, geos=3, scen=3, env=1 → combined=7632
#   Low Priority Pages: pages_each=1526,   scans=6106
#   buffer_scans=1145   predicted=27704    price=4539.68
# GILEAD_env2: geos=3, scen=3, env=2 → combined=15264
#   predicted=55408  price=9028.96
# GILEAD_env1.5: geos=1, scen=3, env=1.5 → combined=3816
#   predicted=13852
# ---------------------------------------------------------------------------

_CADENCE_LAYERS = [
    {"name": "Baseline inventory",      "pct": 1.0,   "runs_per_year": 1,
     "why": "A full sweep so nothing on the site is invisible."},
    {"name": "High Priority",           "pct": 0.015, "runs_per_year": 52,
     "why": "Crown-jewel pages checked weekly."},
    {"name": "Moderate Priority Pages", "pct": 0.075, "runs_per_year": 12,
     "why": "Key templates re-checked monthly."},
    {"name": "Low Priority Pages",      "pct": 0.20,  "runs_per_year": 4,
     "why": "Broad cross-section kept current quarterly."},
]

DATA = {
    "customer": "Arthur J. Gallagher",
    "prepared_by": "Jarrod Wilbur",
    "date": "2026-06-07",
    "use_case": "privacy / consent monitoring",
    "domains": ["ajg.com"],
    "properties_note": "233 web properties identified.",
    "regulations": ["CCPA/CPRA"],
    "monitoring_summary": "Annual full-site privacy & consent monitoring across all properties, "
                          "with quarterly and monthly re-checks of consent-critical pages.",
    "page_count": {"low": 95_721, "anchor": 95_721, "high": 108_155, "confidence": "MEDIUM",
                   "url_total": 485_096, "defensible": 95_721, "discounted": 389_375,
                   "census_id": 711, "crawl_status": "running",
                   "spiral_note": "Excluded ~389,375 query-string-duplicate URLs across 6 properties."},
    "consent_states": {"count": 3, "names": ["Default", "Opt-Out", "GPC"]},
    "multipliers": {"geographies": 1, "scenarios": 3, "environments": 1},
    "cadence_layers": _CADENCE_LAYERS,
    "buffer_pct": 0.15,
    # combined_pages = 95721 × 1 × 3 × 1 = 287163
    # predicted_scans from engine = 1042402
    "usage": {"combined_pages": 287_163, "predicted_scans": 1_042_402},
    "pricing": {
        "predicted_price": 97_984.12,
        "range_low_price": 90_000,
        "range_high_price": 110_000,
        "price_by_band": [
            {"band_limit": 1000,    "rate": 0.0,  "pages": 1000,   "cost": 0.0},
            {"band_limit": 50000,   "rate": 0.17, "pages": 50000,  "cost": 8500.0},
            {"band_limit": 500000,  "rate": 0.12, "pages": 500000, "cost": 60000.0},
            {"band_limit": 1000000, "rate": 0.06, "pages": 491402, "cost": 29484.12},
        ],
        "pricing_source": "live @ https://app.observepoint.com/www-pricing/main.js",
    },
    "internal": {
        "assumptions": ["Geographies defaulted to 1 — confirm regions.",
                        "Consent states assumed CCPA (3) — confirm regulations."],
        "implied_blended_frequency": 3.629,
        "thresholds_swept": "5000/1.3=95721; 10000/1.5=95721; 20000/2.0=106638",
    },
}


def _text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def test_customer_sections_and_derivation():
    t = _text(bp.build_proposal(DATA))
    assert "Arthur J. Gallagher" in t
    assert "Your web footprint" in t
    assert "96,000" in t                              # anchor rounded for customer (not 95,721)
    assert "How your annual usage is calculated" in t
    assert "page scan" in t.lower()                   # the unit is defined
    assert "Default" in t and "Opt-Out" in t and "GPC" in t
    assert "287,163" in t                             # combined pages monitored
    # Four cadence-layer frequencies
    assert "Annually" in t                            # Baseline inventory (runs=1)
    assert "Weekly" in t                              # High Priority (runs=52)
    assert "Monthly" in t                             # Moderate Priority (runs=12)
    assert "Quarterly" in t                           # Low Priority (runs=4)
    assert "1,042,402" in t                           # predicted total
    # Section 4 uses "Annual investment", NOT "Recommended contract"
    assert "Annual investment" in t
    assert "Recommended contract" not in t
    # Exact price = graduated_price(predicted_scans)
    expected_price = cs.graduated_price(DATA["usage"]["predicted_scans"], cs.BAKED_TIERS)["total"]
    assert bp._usd(expected_price) in t               # e.g. "$97,984"
    assert "Scope of Work" in t
    assert "live" in t.lower()


def test_price_is_exact_graduated_price_of_predicted_scans():
    # The price shown MUST match the graduated calculator (the rep's whole ask).
    s = DATA["usage"]["predicted_scans"]
    assert abs(cs.graduated_price(s, cs.BAKED_TIERS)["total"] - DATA["pricing"]["predicted_price"]) < 1


def test_no_internal_section_or_confidence_in_customer_doc():
    t = _text(bp.build_proposal(DATA))
    assert "[INTERNAL" not in t                 # the strippable section is gone
    assert "485,096" not in t                   # raw URL total never in the customer doc
    assert "711" not in t                       # census id never in the customer doc
    assert "CONFIDENCE" not in t.upper()        # confidence moved to the internal file
    for term in ("census", "spiral", "raw url", "defensible", "anchor"):
        assert term not in t.lower(), f"leaked internal term: {term}"


def test_cadence_table_shows_the_why():
    t = _text(bp.build_proposal(DATA))
    assert "Why" in t                                            # the new column header
    assert "nothing on the site is invisible" in t              # a rationale line, customer-facing


def test_clean_guard_rejects_internal_terms_in_narrative():
    d = json.loads(json.dumps(DATA))
    d["monitoring_summary"] = "We discounted query-param spiral URLs."
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_clean_guard_allows_internal_section_terms():
    # 'spiral'/'query-string' etc. live in the INTERNAL fields by design — must NOT trip the guard.
    d = json.loads(json.dumps(DATA))
    d["page_count"]["spiral_note"] = "6 query-param spiral domains discounted."
    bp.build_proposal(d)  # must not raise


def test_clean_guard_allows_collision_identity():
    d = json.loads(json.dumps(DATA))
    d["customer"] = "Discount Tire Co"
    d["domains"] = ["spiral-galaxy.com"]
    d["monitoring_summary"] = "Full-site privacy sweep annually."
    t = _text(bp.build_proposal(d))  # must not raise
    assert "Discount Tire Co" in t


def _with_anchor(low, anchor, high):
    d = json.loads(json.dumps(DATA))
    d["page_count"].update(low=low, anchor=anchor, high=high)
    # Keep the sweep chain consistent with the mutated anchor so the §3 reconciliation guard
    # doesn't trip — these tests exercise §1 footprint rounding, not §3.
    mx = d["multipliers"]
    combined = round(anchor * mx["geographies"] * mx["scenarios"] * mx["environments"])
    d["usage"]["combined_pages"] = combined
    # Recompute predicted_scans from the engine so the predicted-scan guard also passes.
    out = cs.compute({
        "page_count": {"low": anchor, "anchor": anchor, "high": anchor},
        "multipliers": mx,
        "buffer_pct": d["buffer_pct"],
        "tiers": cs.BAKED_TIERS,
        "cadence_layers": d["cadence_layers"],
    })
    d["usage"]["predicted_scans"] = out["anchor"]["predicted_scans"]
    return d


def test_round_sig_preserves_small_and_distinguishes_neighbors():
    # The bug: round-to-nearest-1000 turned ~80 pages into "0" and collided 4,722 & 5,398 at "5,000".
    assert bp._round_sig(80) == 80          # was 0
    assert bp._round_sig(4_722) == 4_700    # was 5,000
    assert bp._round_sig(5_398) == 5_400    # was 5,000 (collided with 4,722)
    assert bp._round_sig(95_721) == 96_000  # matches the methodology's 2-sig-fig example
    assert bp._round_sig(108_155) == 110_000
    assert bp._round_sig(0) == 0


def test_small_footprint_not_zero():
    t = _text(bp.build_proposal(_with_anchor(80, 80, 80)))
    assert "80 pages" in t                       # the real count, not collapsed
    assert "approximately 0 pages" not in t      # the bug symptom: round-to-1000 made ~80 read as 0
    assert "range 0–0" not in t


def test_neighbor_counts_render_distinctly():
    tc = _text(bp.build_proposal(_with_anchor(4_722, 4_722, 4_722)))
    tg = _text(bp.build_proposal(_with_anchor(5_398, 5_398, 5_398)))
    assert "4,700 pages" in tc and "5,400 pages" in tg   # not both collapsed to "5,000 pages"


def test_cli_writes_docx(tmp_path):
    f = tmp_path / "in.json"; f.write_text(json.dumps(DATA))
    out = tmp_path / "p.docx"
    res = subprocess.run([sys.executable, str(SCRIPT), str(f), str(out)],
                         capture_output=True, text=True)
    assert res.returncode == 0, res.stderr
    assert res.stdout.strip() == str(out)
    assert "Arthur J. Gallagher" in _text(Document(out))


# --- friendly input validation (Part B) ------------------------------------------------
def test_cli_friendly_error_missing_required_key(tmp_path):
    bad = json.loads(json.dumps(DATA))
    bad.pop("page_count")
    f = tmp_path / "bad.json"; f.write_text(json.dumps(bad))
    out = tmp_path / "p.docx"
    res = subprocess.run([sys.executable, str(SCRIPT), str(f), str(out)],
                         capture_output=True, text=True)
    assert res.returncode != 0
    assert "missing" in res.stderr.lower() and "page_count" in res.stderr
    assert "Traceback" not in res.stderr and "KeyError" not in res.stderr


def test_cli_friendly_error_missing_usage_keys(tmp_path):
    bad = json.loads(json.dumps(DATA)); bad["usage"] = {"predicted_scans": 1}   # missing combined_pages
    f = tmp_path / "bad.json"; f.write_text(json.dumps(bad)); out = tmp_path / "p.docx"
    res = subprocess.run([sys.executable, str(SCRIPT), str(f), str(out)], capture_output=True, text=True)
    assert res.returncode != 0
    assert "combined_pages" in res.stderr
    assert "Traceback" not in res.stderr
    assert not out.exists()


def test_cli_friendly_error_malformed_json(tmp_path):
    f = tmp_path / "bad.json"; f.write_text("{not valid json")
    out = tmp_path / "p.docx"
    res = subprocess.run([sys.executable, str(SCRIPT), str(f), str(out)],
                         capture_output=True, text=True)
    assert res.returncode != 0
    assert "Traceback" not in res.stderr
    assert "scope-calculator" in res.stderr


def test_clean_guard_rejects_internal_term_in_cadence_name():
    d = json.loads(json.dumps(DATA))
    d["cadence_layers"][1]["name"] = "Quarterly spiral re-check"   # internal term in a customer label
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_clean_guard_rejects_internal_term_in_why():
    d = json.loads(json.dumps(DATA))
    d["cadence_layers"][0]["why"] = "Full crawl to find every defensible page."
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_clean_guard_allows_identity_collision_in_name_and_domain():
    d = json.loads(json.dumps(DATA))
    d["customer"] = "Discount Tire Co"
    d["domains"] = ["spiral-galaxy.com"]
    d["monitoring_summary"] = "Full-site privacy monitoring annually."
    bp.build_proposal(d)  # identity fields are not scrubbed → must not raise


def test_clean_guard_rejects_internal_term_in_properties_note():
    d = json.loads(json.dumps(DATA))
    d["properties_note"] = "Site census crawl excluded spiral pages."
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_proposal_does_not_mention_removed_methodology_sheet_or_internal_terms():
    """Regression: customer workbook has no Methodology sheet; proposal must not advertise it.
    Also guards against internal-derivation language leaking into the rendered customer document."""
    t = _text(bp.build_proposal(DATA))
    # The Methodology sheet was removed from the customer workbook in Phase A.
    assert "Methodology" not in t, "proposal advertises the removed Methodology workbook sheet"
    # Internal-derivation terms must not appear in the customer-facing proposal.
    for term in ("census", "spiral", "raw url", "defensible", "reduced",
                 "crawl", "query-param", "query-string", "recursion", "collapsed"):
        assert term not in t.lower(), f"internal term leaked into proposal: {term!r}"


def test_frequency_advisor_reference_exists_and_has_ladder():
    refs = pathlib.Path(__file__).resolve().parent.parent / "skills" / "scope-calculator" / "references"
    doc = (refs / "frequency-advisor.md").read_text().lower()
    for layer in ("baseline inventory", "high priority", "moderate priority", "low priority", "buffer"):
        assert layer in doc, f"layer missing from frequency-advisor.md: {layer!r}"
    for pct in ("100%", "1.5%", "7.5%", "20%", "15%"):
        assert pct in doc, f"pct missing from frequency-advisor.md: {pct!r}"


# ---------------------------------------------------------------------------
# Gilead-like fixture: multi-geography, full multiplier chain
# combined = 848 × 3 × 3 × 1 = 7,632
# predicted = 27,704   price = 4,539.68   buffer_scans = 1,145
# Low Priority Pages: pages_each=1526, scans=6106
# ---------------------------------------------------------------------------
GILEAD_DATA = {
    "customer": "Gilead Sciences",
    "prepared_by": "Jarrod Wilbur",
    "date": "2026-06-16",
    "use_case": "privacy / consent monitoring",
    "domains": ["gilead.com"],
    "monitoring_summary": "Annual full-site privacy monitoring across all regions and consent states.",
    "page_count": {"low": 700, "anchor": 848, "high": 1000},
    "consent_states": {"count": 3, "names": ["Pre-consent", "Opt-Out", "GPC"]},
    "multipliers": {"geographies": 3, "scenarios": 3, "environments": 1},
    "cadence_layers": _CADENCE_LAYERS,
    "buffer_pct": 0.15,
    # combined = 848 × 3 × 3 × 1 = 7,632
    # predicted from engine = 27,704
    "usage": {"combined_pages": 7_632, "predicted_scans": 27_704},
    "pricing": {
        "predicted_price": 4_539.68,
        "range_low_price": 4_000,
        "range_high_price": 5_500,
        "pricing_source": "live @ https://app.observepoint.com/www-pricing/main.js",
    },
    "regulations": ["CCPA/CPRA"],
}


def test_gilead_sweep_table_has_geographies_row():
    """(a) A Geographies row with ×3 must appear; (b) consent states ×3 present;
    (c) 7,632 is shown; (d) chain reconciles 848×3×3=7,632; (e+f) cadence rows derive
    correctly from pct — not from missing pages/runs keys."""
    t = _text(bp.build_proposal(GILEAD_DATA))

    # (a) geographies row with ×3
    assert "Geographies" in t, "Expected a Geographies row in the sweep table"
    assert "×3" in t

    # (b) consent states ×3
    assert "Pre-consent" in t and "Opt-Out" in t and "GPC" in t
    # consent states row also has ×3
    assert t.count("×3") >= 2  # at least geographies row + consent row

    # (c) pages per full sweep shown
    assert "7,632" in t

    # (d) chain reconciles: 848 × 3 × 3 = 7,632 — anchor must appear
    assert "848" in t

    # (e) Low Priority Pages: pct=0.20, runs=4 → pages_each=round(7632*0.20)=1526; scans=round(7632*0.20*4)=6106
    assert "1,526" in t
    assert "6,106" in t

    # (f) Buffer: round(7632 * 0.15) = 1145; predicted total = 27,704
    assert "1,145" in t
    assert "27,704" in t


def test_gilead_environments_row_hidden_when_one():
    """environments=1 → no Environments row should appear."""
    t = _text(bp.build_proposal(GILEAD_DATA))
    assert "Environments" not in t


def test_gilead_environments_row_shown_when_gt_one():
    """environments=2 → an Environments row must appear.
    GILEAD_env2: combined=15264, predicted=55408, price=9028.96"""
    d = copy.deepcopy(GILEAD_DATA)
    d["multipliers"]["environments"] = 2
    # combined = 848 × 3 × 3 × 2 = 15264; predicted from engine = 55408
    d["usage"]["combined_pages"] = 15_264
    d["usage"]["predicted_scans"] = 55_408
    d["pricing"]["predicted_price"] = 9_028.96
    t = _text(bp.build_proposal(d))
    assert "Environments" in t
    assert "×2" in t


# ---------------------------------------------------------------------------
# Reconciliation guard tests — the adapted _sweep_reconcile_error guard must
# reject payloads where the multiplier chain or predicted total is inconsistent.
# ---------------------------------------------------------------------------

def test_reconcile_guard_rejects_combined_mismatch():
    """combined_pages=848 (anchor only, no multipliers) while multipliers imply 7,632 → ValueError."""
    d = copy.deepcopy(GILEAD_DATA)
    d["usage"]["combined_pages"] = 848   # wrong: multipliers imply 7,632
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_reconcile_guard_rejects_predicted_mismatch():
    """predicted_scans set to half the real total while combined_pages is correct → ValueError."""
    d = copy.deepcopy(GILEAD_DATA)
    d["usage"]["predicted_scans"] = round(27_704 / 2)   # half: wrong total
    with pytest.raises(ValueError):
        bp.build_proposal(d)


def test_consistent_payload_still_renders():
    """The correct GILEAD payload (geos=3, combined=7,632) must build cleanly."""
    bp.build_proposal(GILEAD_DATA)   # must not raise


def test_environments_one_point_five_reconciles():
    """Fractional environment multiplier (1.5) must not trip the reconciliation guard.
    GILEAD_env1.5: geos=1, scen=3, env=1.5 → combined=3816, predicted=13852"""
    d = copy.deepcopy(GILEAD_DATA)
    d["multipliers"] = {"geographies": 1, "scenarios": 3, "environments": 1.5}
    d["usage"]["combined_pages"] = round(848 * 1 * 3 * 1.5)   # 3816
    d["usage"]["predicted_scans"] = 13_852
    d["pricing"]["predicted_price"] = 2_184.84
    bp.build_proposal(d)   # must not raise
